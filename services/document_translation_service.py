"""
Document Translation Service
=============================
Orchestrates the full pipeline:

    Input file (PDF / DOCX)
        │
        ▼
    Extract text   ◄── reuses core.pdf_tools / core.ocr_engine / docx parser
        │
        ▼
    Translate paragraphs  ◄── translation.translator.Translator
        │
        ▼
    Rebuild DOCX (and optionally PDF)
        │
        ▼
    Output file

Does NOT duplicate any existing logic – it imports from core.* modules.
"""
from __future__ import annotations

import os
from pathlib import Path
from typing import Callable, List, Optional, Tuple

from docx import Document as DocxDocument
from docx.shared import Pt, RGBColor
import fitz

import config
from core.ocr_engine import perform_ocr
from core.pdf_tools import is_scanned_pdf, extract_text_from_pdf
from translation.translator import Translator
from utils.file_utils import build_output_path, unique_path
from utils.logger import get_logger

logger = get_logger("document_translation")

ProgressCB = Optional[Callable[[int], None]]
StatusCB   = Optional[Callable[[str], None]]


# ---------------------------------------------------------------------------
# Public entry point
# ---------------------------------------------------------------------------

def translate_document(
    input_path: str,
    target_lang: str,
    source_lang: str = "auto",
    output_path: Optional[str] = None,
    provider: str = config.TRANSLATION_PROVIDER,
    export_pdf: bool = False,
    ocr_dpi: int = 300,
    progress_callback: ProgressCB = None,
    status_callback:   StatusCB   = None,
    **_,
) -> str:
    """
    Translate *input_path* (PDF or DOCX) to *target_lang*.

    Parameters
    ----------
    input_path   : Source document (PDF or DOCX).
    target_lang  : Target language code: "en", "ro", "hu".
    source_lang  : Source language code or "auto".
    output_path  : Destination DOCX path.  Auto-derived if None.
    provider     : Translation backend – "libretranslate", "deepl", "marianmt".
    export_pdf   : If True, also convert the output DOCX to PDF via docx2pdf.
    ocr_dpi      : DPI used when OCR-ing scanned PDFs.

    Returns the output file path (DOCX, or PDF if export_pdf=True).
    """
    _log(status_callback, "Analysing document…")

    ext = Path(input_path).suffix.lower()
    translator = Translator(provider=provider)

    if output_path is None:
        suffix = f"_{target_lang}"
        output_path = unique_path(build_output_path(input_path, None, ".docx", suffix))

    # Step 1 – Extract paragraphs
    if ext == ".pdf":
        paragraphs = _extract_paragraphs_from_pdf(
            input_path, ocr_dpi, status_callback
        )
    elif ext in (".docx", ".doc"):
        paragraphs = _extract_paragraphs_from_docx(input_path, status_callback)
    else:
        raise ValueError(f"Unsupported file type: {ext!r}. Only PDF and DOCX are supported.")

    if not any(p.strip() for p in paragraphs):
        raise RuntimeError("No text found in the document. Nothing to translate.")

    _log(status_callback, f"Extracted {len(paragraphs)} paragraph(s).")

    # Step 2 – Translate
    _log(status_callback, f"Translating using {provider}…")
    _progress(progress_callback, 20)

    def _trans_progress(p: int) -> None:
        # Map 0-100 range of translation into 20-90 range overall
        _progress(progress_callback, 20 + int(p * 0.7))

    translated = translator.translate_paragraphs(
        paragraphs,
        source=source_lang,
        target=target_lang,
        progress_callback=_trans_progress,
        status_callback=status_callback,
    )

    # Step 3 – Rebuild DOCX
    _log(status_callback, "Rebuilding document…")
    _progress(progress_callback, 92)
    _build_docx(translated, output_path, ext, input_path)
    _progress(progress_callback, 97)

    result = output_path

    # Optional: export to PDF
    if export_pdf:
        _log(status_callback, "Exporting PDF…")
        from core.converter import docx_to_pdf
        pdf_out = build_output_path(output_path, None, ".pdf")
        pdf_out = unique_path(pdf_out)
        docx_to_pdf(output_path, pdf_out)
        result = pdf_out

    _progress(progress_callback, 100)
    _log(status_callback, f"Done → {Path(result).name}")
    logger.info("translate_document -> %s", result)
    return result


# ---------------------------------------------------------------------------
# Extraction helpers  (reuse core modules)
# ---------------------------------------------------------------------------

def _extract_paragraphs_from_pdf(
    path: str,
    dpi: int,
    status_callback: StatusCB,
) -> List[str]:
    """
    Extract paragraphs from a PDF.
    - Text-based: use PyMuPDF block extraction (preserves paragraph breaks).
    - Scanned:    run OCR via the existing ocr_engine module.
    """
    _log(status_callback, "Checking if PDF is scanned…")

    if is_scanned_pdf(path):
        _log(status_callback, "Scanned PDF detected – running OCR…")
        import tempfile
        tmp_docx = os.path.join(tempfile.gettempdir(), "_ocr_tmp.docx")
        perform_ocr(path, language="eng", output_path=tmp_docx, dpi=dpi,
                    status_callback=status_callback)
        # Extract from the OCR-generated DOCX
        return _extract_paragraphs_from_docx(tmp_docx, status_callback)

    _log(status_callback, "Extracting text from PDF…")
    paragraphs: List[str] = []

    with fitz.open(path) as doc:
        for page in doc:
            blocks = page.get_text("blocks")   # list of (x0, y0, x1, y1, text, …)
            for block in blocks:
                text = block[4].strip()
                if text:
                    # Each block ≈ a paragraph; split on double newlines within block
                    for para in text.split("\n\n"):
                        clean = para.replace("\n", " ").strip()
                        if clean:
                            paragraphs.append(clean)
            paragraphs.append("")   # blank line between pages

    return paragraphs


def _extract_paragraphs_from_docx(path: str, status_callback: StatusCB) -> List[str]:
    """Extract paragraph texts from a DOCX, preserving empty lines."""
    _log(status_callback, "Extracting paragraphs from DOCX…")
    doc = DocxDocument(path)
    return [p.text for p in doc.paragraphs]   # preserves blanks as ""


# ---------------------------------------------------------------------------
# DOCX reconstruction
# ---------------------------------------------------------------------------

def _build_docx(
    paragraphs: List[str],
    output_path: str,
    source_ext: str,
    source_path: str,
) -> None:
    """
    Build a new DOCX from translated paragraphs.

    For DOCX sources, attempts to mirror the original paragraph styles.
    For PDF sources, uses a clean default style.
    """
    out_doc = DocxDocument()

    if source_ext in (".docx", ".doc"):
        # Mirror original styles
        src_doc = DocxDocument(source_path)
        src_paras = src_doc.paragraphs

        for i, text in enumerate(paragraphs):
            if i < len(src_paras):
                style_name = src_paras[i].style.name
                try:
                    para = out_doc.add_paragraph(text, style=style_name)
                except Exception:
                    para = out_doc.add_paragraph(text)
            else:
                out_doc.add_paragraph(text)
    else:
        # PDF source – plain output with readable formatting
        for text in paragraphs:
            if text.strip():
                out_doc.add_paragraph(text)
            else:
                out_doc.add_paragraph("")

    out_doc.save(output_path)


# ---------------------------------------------------------------------------
# Utility
# ---------------------------------------------------------------------------

def _progress(cb: ProgressCB, value: int) -> None:
    if cb:
        cb(value)


def _log(cb: StatusCB, msg: str) -> None:
    if cb:
        cb(msg)
    logger.debug(msg)
