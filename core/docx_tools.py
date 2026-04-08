"""
DOCX manipulation utilities.
"""

from __future__ import annotations

import os
from pathlib import Path
from typing import Callable, List, Optional

from docx import Document
from PIL import Image

from utils.file_utils import unique_path, ensure_temp_dir
from utils.logger import get_logger

logger = get_logger("docx_tools")

ProgressCB = Optional[Callable[[int], None]]
StatusCB = Optional[Callable[[str], None]]


def _report(cb: ProgressCB, value: int) -> None:
    if cb:
        cb(value)


def _status(cb: StatusCB, msg: str) -> None:
    if cb:
        cb(msg)
    logger.debug(msg)


# ---------------------------------------------------------------------------
# Extract images
# ---------------------------------------------------------------------------

def extract_images_from_docx(
    input_path: str,
    output_dir: str,
    progress_callback: ProgressCB = None,
    status_callback: StatusCB = None,
    **_,
) -> List[str]:
    """Save every embedded image in the DOCX to *output_dir*."""
    os.makedirs(output_dir, exist_ok=True)
    saved: List[str] = []

    doc = Document(input_path)
    rels = doc.part.rels
    total = len(rels)

    image_rels = [
        r for r in rels.values()
        if "image" in r.target_ref
    ]

    for idx, rel in enumerate(image_rels):
        _status(status_callback, f"Extracting image {idx + 1}/{len(image_rels)}…")
        try:
            blob = rel.target_part.blob
            ext = Path(rel.target_ref).suffix or ".png"
            out_path = unique_path(
                os.path.join(output_dir, f"image_{idx + 1}{ext}")
            )
            with open(out_path, "wb") as f:
                f.write(blob)
            saved.append(out_path)
        except Exception as exc:
            logger.warning("Could not extract image %d: %s", idx + 1, exc)
        _report(progress_callback, int((idx + 1) / max(len(image_rels), 1) * 100))

    logger.info("extract_images_from_docx -> %d images", len(saved))
    return saved


# ---------------------------------------------------------------------------
# DOCX → images
# ---------------------------------------------------------------------------

def docx_to_images(
    input_path: str,
    output_dir: str,
    dpi: int = 150,
    progress_callback: ProgressCB = None,
    status_callback: StatusCB = None,
    **_,
) -> List[str]:
    """
    Convert each page of a DOCX to a PNG image.

    Requires LibreOffice to be installed (used to render the DOCX to PDF first,
    then PyMuPDF converts each page to an image).
    """
    import subprocess
    import fitz

    os.makedirs(output_dir, exist_ok=True)
    stem = Path(input_path).stem
    saved: List[str] = []

    # Step 1 – convert DOCX to PDF via LibreOffice
    _status(status_callback, "Converting DOCX to PDF via LibreOffice…")
    tmp_dir = ensure_temp_dir()
    cmd = _libreoffice_cmd(input_path, tmp_dir)
    if cmd is None:
        raise RuntimeError(
            "LibreOffice not found. Please install LibreOffice to use this feature."
        )

    result = subprocess.run(cmd, capture_output=True, text=True)
    if result.returncode != 0:
        raise RuntimeError(f"LibreOffice error:\n{result.stderr}")

    pdf_path = os.path.join(tmp_dir, f"{stem}.pdf")
    if not os.path.isfile(pdf_path):
        raise RuntimeError("LibreOffice did not produce a PDF output file.")

    # Step 2 – render each PDF page to PNG
    _status(status_callback, "Rendering pages to images…")
    with fitz.open(pdf_path) as doc:
        total = doc.page_count
        scale = dpi / 72.0
        mat = fitz.Matrix(scale, scale)
        for i, page in enumerate(doc):
            pix = page.get_pixmap(matrix=mat, alpha=False)
            out_path = unique_path(os.path.join(output_dir, f"{stem}_page{i + 1}.png"))
            pix.save(out_path)
            saved.append(out_path)
            _report(progress_callback, int((i + 1) / total * 100))

    logger.info("docx_to_images -> %d images", len(saved))
    return saved


# ---------------------------------------------------------------------------
# Document info
# ---------------------------------------------------------------------------

def get_docx_info(input_path: str, **_) -> dict:
    """Return basic metadata about the DOCX."""
    doc = Document(input_path)
    props = doc.core_properties
    para_count = len(doc.paragraphs)
    word_count = sum(len(p.text.split()) for p in doc.paragraphs)
    return {
        "title": props.title or "",
        "author": props.author or "",
        "created": str(props.created or ""),
        "modified": str(props.modified or ""),
        "paragraphs": para_count,
        "words": word_count,
    }


# ---------------------------------------------------------------------------
# Internal helpers
# ---------------------------------------------------------------------------

def _libreoffice_cmd(input_path: str, output_dir: str) -> Optional[list]:
    import config
    for lo_path in config.LIBREOFFICE_PATHS:
        if os.path.isfile(lo_path):
            return [
                lo_path,
                "--headless",
                "--convert-to", "pdf",
                "--outdir", output_dir,
                input_path,
            ]
    # Try PATH
    import shutil as sh
    if sh.which("soffice"):
        return [
            "soffice",
            "--headless",
            "--convert-to", "pdf",
            "--outdir", output_dir,
            input_path,
        ]
    return None
